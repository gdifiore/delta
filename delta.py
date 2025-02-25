from docx import Document
import os
import json
import hashlib
from datetime import datetime
from dataclasses import dataclass, asdict
from typing import Dict, List, Optional, Tuple
import difflib
from pathlib import Path
import gzip
import re
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from dataclasses import dataclass, asdict, field


@dataclass
class RichContent:
    """Base class for rich content elements in a document."""
    type: str  # 'paragraph', 'table', 'image'


@dataclass
class ParagraphContent(RichContent):
    type: str = 'paragraph'
    text: str = ''
    bold: bool = False
    italic: bool = False


@dataclass
class TableContent(RichContent):
    type: str = 'table'
    # Simplified: list of rows, each a list of cell texts
    rows: List[List[str]] = field(default_factory=list)


@dataclass
class ImageContent(RichContent):
    type: str = 'image'
    id: str = ''  # Unique identifier from the document
    data: bytes = b''  # Binary image data


@dataclass
class DeltaOperation:
    """Represents a single incremental change operation on rich content."""
    op: str  # 'insert', 'delete', 'replace'
    position: int  # Position in the content list
    # Rich content object or None for delete
    content: Optional[RichContent] = None


@dataclass
class Delta:
    timestamp: str
    parent_hash: str
    baseline_hash: str
    operations: List[DeltaOperation]  # Operations on rich content
    message: str
    hash: str
    is_snapshot: bool = False  # Flag to indicate full content storage

    def to_dict(self):
        """Convert Delta to a JSON-serializable dictionary."""
        return {
            "timestamp": self.timestamp,
            "parent_hash": self.parent_hash,
            "baseline_hash": self.baseline_hash,
            "operations": [
                {
                    "op": op.op,
                    "position": op.position,
                    "content": self._serialize_content(op.content) if op.content else None
                }
                for op in self.operations
            ],
            "message": self.message,
            "hash": self.hash,
            "is_snapshot": self.is_snapshot
        }

    def _serialize_content(self, content):
        """Helper method to serialize RichContent objects safely."""
        if isinstance(content, ParagraphContent):
            return {
                "type": "paragraph",
                "text": content.text,
                "bold": content.bold,
                "italic": content.italic
            }
        elif isinstance(content, TableContent):
            return {
                "type": "table",
                "rows": content.rows
            }
        elif isinstance(content, ImageContent):
            return {
                "type": "image",
                "id": content.id,
                "data": content.data.hex()  # Convert binary image data to hex string
            }
        return None  # Catch-all for unexpected content types


class CombinedDiffer:
    def __init__(self):
        self.COLORS = {
            'GREEN': '\033[92m', 'RED': '\033[91m', 'BLUE': '\033[94m',
            'YELLOW': '\033[93m', 'GREY': '\033[90m', 'RESET': '\033[0m'
        }

    def _split_into_words(self, text: str) -> List[str]:
        """Split text into words while preserving whitespace and punctuation."""
        return re.findall(r'\S+|\s+', text)

    def _normalize_text(self, text: str) -> str:
        """Normalize text for comparison (strip whitespace, lowercase)."""
        return re.sub(r'\s+', ' ', text.strip().lower())

    def _find_moved_blocks(self, old_paragraphs: List[str], new_paragraphs: List[str],
                           min_block_length: int = 5) -> Dict[str, 'TextBlock']:
        """Identify blocks of text that have been moved."""
        from dataclasses import dataclass

        @dataclass
        class TextBlock:
            content: str
            old_position: int = -1
            new_position: int = -1

        blocks = {}
        old_normalized = [self._normalize_text(p) for p in old_paragraphs]
        new_normalized = [self._normalize_text(p) for p in new_paragraphs]

        for old_idx, old_text in enumerate(old_paragraphs):
            old_words = old_text.split()
            if len(old_words) < min_block_length:
                continue
            old_norm = old_normalized[old_idx]
            if old_norm in blocks:
                continue
            for new_idx, new_norm in enumerate(new_normalized):
                if old_norm == new_norm and old_idx != new_idx:
                    blocks[old_norm] = TextBlock(
                        content=old_text,
                        old_position=old_idx,
                        new_position=new_idx
                    )
                    break

        return blocks

    def format_word_diff(self, old_text: str, new_text: str) -> str:
        """Generate a formatted word-level diff."""
        old_words = self._split_into_words(old_text)
        new_words = self._split_into_words(new_text)
        matcher = difflib.SequenceMatcher(None, old_words, new_words)
        result = []
        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                result.append(''.join(old_words[i1:i2]))
            elif op == 'delete':
                result.append(self._color_text(
                    ''.join(old_words[i1:i2]), 'RED'))
            elif op == 'insert':
                result.append(self._color_text(
                    ''.join(new_words[j1:j2]), 'GREEN'))
            elif op == 'replace':
                result.append(self._color_text(
                    ''.join(old_words[i1:i2]), 'RED'))
                result.append(self._color_text(
                    ''.join(new_words[j1:j2]), 'GREEN'))
        return ''.join(result)

    def format_combined_diff(self, old_content: List[RichContent], new_content: List[RichContent]) -> str:
        """Format a terminal-friendly diff showing rich content changes."""
        result = []
        matcher = difflib.SequenceMatcher(None, [str(c) for c in old_content], [
                                          str(c) for c in new_content])

        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                for idx in range(i1, i2):
                    item = old_content[idx]
                    if item.type == 'paragraph':
                        result.append(item.text)
                    # Skip tables/images in unchanged sections for brevity
            elif op == 'delete':
                for idx in range(i1, i2):
                    item = old_content[idx]
                    if item.type == 'paragraph':
                        result.append(self._color_text(
                            f"- {item.text}", 'RED'))
                    elif item.type == 'table':
                        result.append(self._color_text(
                            "[Removed Table]", 'RED'))
                    elif item.type == 'image':
                        result.append(self._color_text(
                            "[Removed Image]", 'RED'))
            elif op == 'insert':
                for idx in range(j1, j2):
                    item = new_content[idx]
                    if item.type == 'paragraph':
                        result.append(self._color_text(
                            f"+ {item.text}", 'GREEN'))
                    elif item.type == 'table':
                        result.append(self._color_text("[New Table]", 'GREEN'))
                    elif item.type == 'image':
                        result.append(self._color_text("[New Image]", 'GREEN'))
            elif op == 'replace':
                for idx in range(i1, i2):
                    item = old_content[idx]
                    if item.type == 'paragraph':
                        result.append(self._color_text(
                            f"- {item.text}", 'RED'))
                    elif item.type == 'table':
                        result.append(self._color_text(
                            "[Removed Table]", 'RED'))
                    elif item.type == 'image':
                        result.append(self._color_text(
                            "[Removed Image]", 'RED'))
                for idx in range(j1, j2):
                    item = new_content[idx]
                    if item.type == 'paragraph':
                        result.append(self._color_text(
                            f"+ {item.text}", 'GREEN'))
                    elif item.type == 'table':
                        result.append(self._color_text("[New Table]", 'GREEN'))
                    elif item.type == 'image':
                        result.append(self._color_text("[New Image]", 'GREEN'))

        return '\n'.join(result)

    def _add_content_to_doc(self, doc: Document, content: RichContent) -> None:
        """Add unchanged content to the DOCX."""
        if content.type == 'paragraph':
            p = doc.add_paragraph(content.text)
            if content.bold:
                p.runs[0].bold = True
            if content.italic:
                p.runs[0].italic = True
        elif content.type == 'table':
            table = doc.add_table(rows=len(content.rows),
                                  cols=len(content.rows[0]))
            for i, row in enumerate(content.rows):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = cell_text
        elif content.type == 'image':
            with io.BytesIO(content.data) as image_stream:
                doc.add_picture(image_stream)

    def _add_deleted_content(self, doc: Document, content: RichContent) -> None:
        """Add deleted content with strikethrough."""
        if content.type == 'paragraph':
            p = doc.add_paragraph()
            run = p.add_run(content.text)
            run.font.strike = True
            if content.bold:
                run.bold = True
            if content.italic:
                run.italic = True
        # Tables and images marked as deleted with a note
        elif content.type == 'table' or content.type == 'image':
            p = doc.add_paragraph(f"[Deleted {content.type}]")
            run = p.runs[0]
            run.font.strike = True

    def _add_inserted_content(self, doc: Document, content: RichContent) -> None:
        """Add inserted content with underline."""
        if content.type == 'paragraph':
            p = doc.add_paragraph()
            run = p.add_run(content.text)
            run.font.underline = True
            if content.bold:
                run.bold = True
            if content.italic:
                run.italic = True
        elif content.type == 'table':
            table = doc.add_table(rows=len(content.rows),
                                  cols=len(content.rows[0]))
            for i, row in enumerate(content.rows):
                for j, cell_text in enumerate(row):
                    table.rows[i].cells[j].text = cell_text
            # Underline table text as a proxy for "inserted"
            for row in table.rows:
                for cell in row.cells:
                    for par in cell.paragraphs:
                        par.runs[0].font.underline = True
        elif content.type == 'image':
            with io.BytesIO(content.data) as image_stream:
                doc.add_picture(image_stream)
            doc.add_paragraph("[Inserted image]").runs[0].font.underline = True

    def _color_text(self, text: str, color: str) -> str:
        """Wrap text in color codes."""
        return f"{self.COLORS[color]}{text}{self.COLORS['RESET']}"


class DocxVersionStore:
    SNAPSHOT_INTERVAL = 10  # Create a full snapshot every 10 commits

    def __init__(self, path: str = ".delta"):
        self.root_path = Path(path)
        self.objects_path = self.root_path / "objects"
        self.refs_path = self.root_path / "refs"
        self.head_path = self.root_path / "HEAD"
        if not self.root_path.exists():
            self.initialize_store()
        else:
            self._ensure_store_structure()

    def _ensure_store_structure(self) -> None:
        """Ensure the store structure is intact or create it if missing."""
        for directory in [self.objects_path, self.refs_path]:
            directory.mkdir(parents=True, exist_ok=True)
        if not self.head_path.exists():
            self.head_path.touch()

    def initialize_store(self) -> None:
        """Initialize a new version store."""
        self.objects_path.mkdir(parents=True)
        self.refs_path.mkdir(parents=True)
        self.head_path.touch()
        (self.refs_path / "master").touch()

    def _extract_content(self, doc_input: str | io.BytesIO) -> List[RichContent]:
        if isinstance(doc_input, str):
            doc = Document(doc_input)
        else:  # Assume BytesIO
            doc = Document(doc_input)
        content = []
        for element in doc.element.body:
            if element.tag.endswith('p'):
                para = doc.paragraphs[len(content)] if len(content) < len(doc.paragraphs) else None
                if para:
                    bold = any(r.font.bold for r in para.runs if r.font.bold is not None)
                    italic = any(r.font.italic for r in para.runs if r.font.italic is not None)
                    content.append(ParagraphContent(text=para.text, bold=bold, italic=italic))
            elif element.tag.endswith('tbl'):
                table = doc.tables[len([c for c in content if c.type == 'table'])]
                rows = [[cell.text for cell in row.cells] for row in table.rows]
                content.append(TableContent(rows=rows))
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                image_part = rel.target_part
                content.append(ImageContent(id=image_part.partname, data=image_part.blob))
        return content

    def _calculate_hash(self, content: List[RichContent]) -> str:
        """Calculate a hash for rich content."""
        hasher = hashlib.sha256()
        for item in content:
            if item.type == 'paragraph':
                hasher.update(
                    f"{item.text}{item.bold}{item.italic}".encode('utf-8'))
            elif item.type == 'table':
                hasher.update(''.join(''.join(cell for cell in row)
                              for row in item.rows).encode('utf-8'))
            elif item.type == 'image':
                hasher.update(item.data)
        return hasher.hexdigest()[:8]

    def _compute_operations(self, old_content: List[RichContent], new_content: List[RichContent]) -> List[DeltaOperation]:
        """Compute incremental operations between two rich content versions."""
        matcher = difflib.SequenceMatcher(None, [str(c) for c in old_content], [
                                          str(c) for c in new_content])
        operations = []

        for op, i1, i2, j1, j2 in matcher.get_opcodes():
            if op == 'equal':
                continue
            elif op == 'delete':
                for pos in range(i1, i2):
                    operations.append(DeltaOperation('delete', pos, None))
            elif op == 'insert':
                for pos, idx in enumerate(range(j1, j2), i1):
                    operations.append(DeltaOperation(
                        'insert', pos, new_content[idx]))
            elif op == 'replace':
                for pos, idx in enumerate(range(j1, j2), i1):
                    operations.append(DeltaOperation(
                        'replace', pos, new_content[idx]))

        return operations

    def _store_object(self, delta: Delta) -> None:
        object_path = self.objects_path / delta.hash
        print(f"Storing object at: {object_path}")
        if delta.is_snapshot:
            snapshot_path = object_path.with_suffix('.docx.gz')
            try:
                with gzip.open(snapshot_path, 'wb') as f:
                    f.write(delta.operations[0].content)
                print(f"Stored snapshot bytes at: {snapshot_path}")
                # Store hex string in JSON
                delta.operations[0].content = delta.operations[0].content.hex()
            except Exception as e:
                print(f"Failed to store snapshot bytes: {e}")
                raise
            json_data = json.dumps(delta.to_dict(), indent=2, ensure_ascii=False)
            try:
                with gzip.open(object_path, 'wt', encoding='utf-8') as f:
                    f.write(json_data)
                print(f"Stored metadata at: {object_path}")
            except Exception as e:
                print(f"Failed to store metadata: {e}")
                raise
        else:
            json_data = json.dumps(delta.to_dict(), indent=2, ensure_ascii=False)
            with gzip.open(object_path, 'wt', encoding='utf-8') as f:
                f.write(json_data)

    def _load_object(self, hash_id: str) -> Optional[Delta]:
        print(f"Loading object for hash_id: {hash_id}")
        object_path = self.objects_path / hash_id
        print(f"Object path: {object_path}")
        if not object_path.exists():
            print(f"Path does not exist: {object_path}")
            return None
        try:
            with gzip.open(object_path, 'rt', encoding='utf-8') as f:
                data = json.load(f)
                print(f"Loaded JSON data: {data}")
                if data.get('is_snapshot', False):
                    snapshot_path = object_path.with_suffix('.docx.gz')
                    if not snapshot_path.exists():
                        print(f"Snapshot file missing: {snapshot_path}")
                        return None
                    with gzip.open(snapshot_path, 'rb') as gz:
                        docx_data = gz.read()
                    return Delta(
                        timestamp=data['timestamp'],
                        parent_hash=data['parent_hash'],
                        baseline_hash=data['baseline_hash'],
                        operations=[DeltaOperation('insert', 0, docx_data)],
                        message=data['message'],
                        hash=data['hash'],
                        is_snapshot=True
                    )
                operations = [DeltaOperation(**op) for op in data['operations']]
                data['operations'] = operations
                return Delta(**data)
        except (gzip.BadGzipFile, json.JSONDecodeError, FileNotFoundError, KeyError) as e:
            print(f"Load error for {hash_id}: {str(e)}")
            raise ValueError(f"Failed to load delta {hash_id}: {str(e)}") from e

    def _update_ref(self, ref_name: str, hash_id: str) -> None:
        """Update a reference to point to a hash."""
        (self.refs_path / ref_name).write_text(hash_id)

    def _get_ref(self, ref_name: str) -> Optional[str]:
        """Get the hash a reference points to."""
        ref_path = self.refs_path / ref_name
        return ref_path.read_text().strip() if ref_path.exists() else None

    def _update_head(self, ref_name: str) -> None:
        """Update HEAD to point to a reference."""
        self.head_path.write_text(f"ref: {ref_name}")

    def _get_head(self) -> Optional[str]:
        if not self.head_path.exists():
            print("HEAD file does not exist")
            return None
        head_content = self.head_path.read_text().strip()
        print(f"HEAD content: {head_content}")
        return self._get_ref(head_content[5:]) if head_content.startswith("ref: ") else head_content

    def _get_content(self, version_hash: str) -> List[RichContent] or bytes: # type: ignore
        print(f"Fetching content for version_hash: {version_hash}")
        delta = self._load_object(version_hash)
        if not delta:
            print(f"Delta not found for hash: {version_hash}")
            raise ValueError(f"Version {version_hash} not found")
        if delta.is_snapshot:
            return delta.operations[0].content
        if not delta.parent_hash:  # Handle empty parent_hash
            content = []  # Start with empty content for the first delta
        else:
            parent_content = self._get_content(delta.parent_hash)
            if isinstance(parent_content, bytes):
                doc = Document(io.BytesIO(parent_content))
                content = self._extract_content(doc)
            else:
                content = parent_content.copy()
        for op in delta.operations:
            if op.op == 'insert':
                content.insert(op.position, op.content)
            elif op.op == 'delete':
                if op.position < len(content):
                    del content[op.position]
            elif op == 'replace':
                if op.position < len(content):
                    content[op.position] = op.content
        return content

    def commit(self, docx_path: str, message: str) -> str:
        with open(docx_path, 'rb') as f:
            docx_bytes = f.read()
        content = self._extract_content(docx_path)
        head = self._get_head()
        print(f"HEAD hash: {head}")
        content_hash = self._calculate_hash(content)
        commit_count = len(self.get_history()) + 1
        print(f"Commit count: {commit_count}")

        is_snapshot = (not head) or (commit_count % self.SNAPSHOT_INTERVAL == 0)
        print(f"is_snapshot: {is_snapshot}")
        if is_snapshot:
            operations = [DeltaOperation('insert', 0, docx_bytes)]
            delta = Delta(
                timestamp=datetime.now().isoformat(),
                parent_hash=head or "",
                baseline_hash=content_hash,
                operations=operations,
                message=message,
                hash=content_hash,
                is_snapshot=True
            )
        else:
            parent_content = self._get_content(head)
            if isinstance(parent_content, bytes):
                parent_content = self._extract_content(io.BytesIO(parent_content))
            elif parent_content is None:
                parent_content = []
            operations = self._compute_operations(parent_content, content)
            parent_delta = self._load_object(head)
            baseline_hash = parent_delta.baseline_hash if parent_delta else content_hash
            delta = Delta(
                timestamp=datetime.now().isoformat(),
                parent_hash=head,
                baseline_hash=baseline_hash,
                operations=operations,
                message=message,
                hash=content_hash
            )
        self._store_object(delta)
        self._update_ref("master", content_hash)
        self._update_head("master")
        return content_hash

    def get_history(self) -> List[Tuple[str, str, datetime]]:
        history = []
        current = self._get_head()
        while current:
            delta = self._load_object(current)
            if delta:
                history.append((delta.hash, delta.message, datetime.fromisoformat(delta.timestamp)))
                current = delta.parent_hash
            else:
                break
        return history

    def export_version(self, version_hash: str, output_path: str) -> None:
        """Export a version to a DOCX file."""
        content = self._get_content(version_hash)

        # Handle raw bytes content (from snapshots)
        if isinstance(content, bytes):
            with open(output_path, 'wb') as f:
                f.write(content)
            return

        # Create new document
        doc = Document()

        for item in content:
            if item.type == 'paragraph':
                # Add paragraph
                p = doc.add_paragraph()
                # Add a single run with the text and formatting
                run = p.add_run(item.text)
                if item.bold:
                    run.font.bold = True
                if item.italic:
                    run.font.italic = True

            elif item.type == 'table':
                # Create table with exact dimensions
                table = doc.add_table(rows=len(item.rows), cols=len(item.rows[0]))

                # Populate table cells
                for i, row in enumerate(item.rows):
                    for j, cell_text in enumerate(row):
                        cell = table.cell(i, j)
                        # Add text directly to the cell
                        cell.text = cell_text

            elif item.type == 'image':
                # Add image directly from binary data
                image_stream = io.BytesIO(item.data)
                doc.add_picture(image_stream)

        # Save the document
        doc.save(output_path)

    def show_diff(self, version_hash: str) -> Optional[List[str]]:
        """Show unified diff for a version (for compatibility)."""
        delta = self._load_object(version_hash)
        if not delta or not delta.parent_hash:
            return None
        parent_content = self._get_content(delta.parent_hash)
        current_content = self._get_content(version_hash)
        return list(difflib.unified_diff(parent_content, current_content, fromfile='previous', tofile='current', lineterm=''))

    def _generate_changes_summary(self, old_content: List[str], new_content: List[str]) -> str:
        """Generate a human-readable summary of changes."""
        old_len = len(old_content)
        new_len = len(new_content)
        added = len([line for line in new_content if line not in old_content])
        removed = len(
            [line for line in old_content if line not in new_content])
        summary_parts = []
        if added:
            summary_parts.append(
                f"{added} line{'s' if added != 1 else ''} added")
        if removed:
            summary_parts.append(
                f"{removed} line{'s' if removed != 1 else ''} removed")
        if not summary_parts:
            return "No text changes" if old_len == new_len else "Content rearranged"
        return "Changes: " + ", ".join(summary_parts)

    def show_friendly_diff(self, version_hash: str) -> str:
        """Print a terminal-friendly diff for a version."""
        delta = self._load_object(version_hash)
        if not delta or not delta.parent_hash:
            return "Initial version"

        parent_content = self._get_content(delta.parent_hash)
        if isinstance(parent_content, bytes):
            doc = Document(io.BytesIO(parent_content))
            parent_content = self._extract_content(doc)
        current_content = self._get_content(version_hash)
        if isinstance(current_content, bytes):
            doc = Document(io.BytesIO(current_content))
            current_content = self._extract_content(doc)

        differ = CombinedDiffer()
        diff_text = differ.format_combined_diff(
            parent_content, current_content)
        summary = self._generate_changes_summary(
            [c.text if c.type == 'paragraph' else '' for c in parent_content],
            [c.text if c.type == 'paragraph' else '' for c in current_content]
        )
        return f"\n=== {summary} ===\n{diff_text}"

    def compare_versions(self, hash1: str, hash2: str) -> List[str]:
        """Compare two versions and return a unified diff."""
        content1 = self._get_content(hash1)
        content2 = self._get_content(hash2)
        return list(difflib.unified_diff(content1, content2, fromfile='version1', tofile='version2', lineterm=''))
